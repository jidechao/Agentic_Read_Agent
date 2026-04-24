"""Multi-format document ingester — parses PDF, Markdown, DOCX, HTML into a unified model."""
import re
from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF
from bs4 import BeautifulSoup
from docx import Document as DocxDocument


# ── Data model ────────────────────────────────────────────────────────────


@dataclass(frozen=True)
class Heading:
    """A single heading extracted from a document."""

    level: int  # 1-6
    text: str


@dataclass
class IngestResult:
    """Unified result from document ingestion."""

    text: str
    title: str
    headings: list[Heading] = field(default_factory=list)
    has_tables: bool = False
    has_lists: bool = False
    metadata: dict = field(default_factory=dict)


# ── Format mapping ────────────────────────────────────────────────────────

SUPPORTED_FORMATS: dict[str, str] = {
    ".pdf": "pdf",
    ".md": "md",
    ".markdown": "md",
    ".docx": "docx",
    ".doc": "docx",
    ".html": "html",
    ".htm": "html",
}


# ── Ingester ──────────────────────────────────────────────────────────────


class DocumentIngester:
    """Parses documents in multiple formats into a unified IngestResult."""

    def ingest(self, file_path: Path) -> IngestResult:
        """Ingest a document file and return a structured result.

        Args:
            file_path: Path to the document file.

        Returns:
            IngestResult with extracted text, title, headings, and metadata.

        Raises:
            ValueError: If the file format is not supported.
            FileNotFoundError: If the file does not exist.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        fmt = self._detect_format(path)
        parser = {
            "pdf": self._parse_pdf,
            "md": self._parse_markdown,
            "docx": self._parse_docx,
            "html": self._parse_html,
        }[fmt]
        return parser(path)

    def _detect_format(self, file_path: Path) -> str:
        """Detect document format from file extension.

        Args:
            file_path: Path to the document.

        Returns:
            Format key string (e.g. "pdf", "md", "docx", "html").

        Raises:
            ValueError: If the file extension is not in SUPPORTED_FORMATS.
        """
        suffix = file_path.suffix.lower()
        fmt = SUPPORTED_FORMATS.get(suffix)
        if fmt is None:
            raise ValueError(f"不支持的格式: {suffix}")
        return fmt

    # ── Markdown ──────────────────────────────────────────────────────────

    def _parse_markdown(self, file_path: Path) -> IngestResult:
        """Parse a Markdown file into IngestResult."""
        content = file_path.read_text(encoding="utf-8")

        heading_pattern = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)
        headings: list[Heading] = [
            Heading(level=len(m.group(1)), text=m.group(2).strip())
            for m in heading_pattern.finditer(content)
        ]

        title = headings[0].text if headings else file_path.stem

        has_tables = "|" in content and "---" in content
        has_lists = bool(re.search(r"^\s*[-*+]\s", content, re.MULTILINE))

        return IngestResult(
            text=content,
            title=title,
            headings=headings,
            has_tables=has_tables,
            has_lists=has_lists,
            metadata={"format": "markdown", "file": str(file_path)},
        )

    # ── PDF ───────────────────────────────────────────────────────────────

    def _parse_pdf(self, file_path: Path) -> IngestResult:
        """Parse a PDF file into IngestResult using PyMuPDF."""
        doc = fitz.open(str(file_path))
        try:
            text_parts: list[str] = []
            headings: list[Heading] = []

            for page in doc:
                page_dict = page.get_text("dict")
                for block in page_dict.get("blocks", []):
                    if block.get("type") != 0:  # text blocks only
                        continue
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            span_text = span.get("text", "").strip()
                            if not span_text:
                                continue
                            text_parts.append(span_text)
                            font_size = span.get("size", 0)
                            if font_size > 14:
                                headings.append(Heading(level=1, text=span_text))

            full_text = "\n".join(text_parts)
            # Pick first heading with actual text content (skip emoji-only)
            title = file_path.stem
            for h in headings:
                if h.text and any(c.isalnum() or '\u4e00' <= c <= '\u9fff' for c in h.text):
                    title = h.text
                    break

            return IngestResult(
                text=full_text,
                title=title,
                headings=headings,
                metadata={"format": "pdf", "file": str(file_path)},
            )
        finally:
            doc.close()

    # ── DOCX ──────────────────────────────────────────────────────────────

    def _parse_docx(self, file_path: Path) -> IngestResult:
        """Parse a DOCX file into IngestResult using python-docx."""
        doc = DocxDocument(str(file_path))

        headings: list[Heading] = []
        text_parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            text_parts.append(text)

            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                level_str = style_name.replace("Heading", "").strip()
                level = int(level_str) if level_str else 1
                headings.append(Heading(level=level, text=text))

        title = headings[0].text if headings else file_path.stem
        has_tables = len(doc.tables) > 0

        full_text = "\n".join(text_parts)

        return IngestResult(
            text=full_text,
            title=title,
            headings=headings,
            has_tables=has_tables,
            metadata={"format": "docx", "file": str(file_path)},
        )

    # ── HTML ──────────────────────────────────────────────────────────────

    def _parse_html(self, file_path: Path) -> IngestResult:
        """Parse an HTML file into IngestResult using BeautifulSoup."""
        content = file_path.read_text(encoding="utf-8")
        soup = BeautifulSoup(content, "lxml")

        headings: list[Heading] = []
        for tag in soup.find_all(re.compile(r"^h[1-6]$")):
            level = int(tag.name[1])
            text = tag.get_text(strip=True)
            if text:
                headings.append(Heading(level=level, text=text))

        title = headings[0].text if headings else file_path.stem

        paragraphs = [p.get_text(strip=True) for p in soup.find_all("p")]
        full_text = "\n".join(paragraphs)

        has_tables = soup.find("table") is not None
        has_lists = soup.find("ul") is not None or soup.find("ol") is not None

        return IngestResult(
            text=full_text,
            title=title,
            headings=headings,
            has_tables=has_tables,
            has_lists=has_lists,
            metadata={"format": "html", "file": str(file_path)},
        )
