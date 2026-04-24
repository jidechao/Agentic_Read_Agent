"""Tests for DocumentIngester — multi-format document parsing (MD/PDF/DOCX/HTML)."""
import os
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from src.ingester import DocumentIngester, Heading, IngestResult


# ── Markdown ──────────────────────────────────────────────────────────────


def test_ingest_markdown(tmp_path: Path) -> None:
    """Markdown with h1 + h2 headings should produce structured IngestResult."""
    md_file = tmp_path / "sample.md"
    md_file.write_text(
        "# Main Title\n\nSome intro text.\n\n"
        "## Section A\n\nContent of section A.\n\n"
        "### Subsection\n\nMore detail.\n",
        encoding="utf-8",
    )

    result = DocumentIngester().ingest(md_file)

    assert isinstance(result, IngestResult)
    assert result.title == "Main Title"
    assert len(result.headings) == 3
    assert result.headings[0] == Heading(level=1, text="Main Title")
    assert result.headings[1] == Heading(level=2, text="Section A")
    assert result.headings[2] == Heading(level=3, text="Subsection")
    assert "Main Title" in result.text
    assert "Section A" in result.text


def test_ingest_markdown_flat(tmp_path: Path) -> None:
    """Markdown without headings should use file stem as title and report no structure."""
    md_file = tmp_path / "flat_doc.md"
    md_file.write_text(
        "Just some plain text without any headings.\n"
        "Another paragraph here.\n",
        encoding="utf-8",
    )

    result = DocumentIngester().ingest(md_file)

    assert result.title == "flat_doc"
    assert result.headings == []
    assert "plain text" in result.text


def test_ingest_markdown_tables_and_lists(tmp_path: Path) -> None:
    """Markdown with tables and lists should set has_tables and has_lists."""
    md_file = tmp_path / "rich.md"
    md_file.write_text(
        "# Rich Doc\n\n"
        "| A | B |\n| --- | --- |\n| 1 | 2 |\n\n"
        "- item one\n- item two\n",
        encoding="utf-8",
    )

    result = DocumentIngester().ingest(md_file)

    assert result.has_tables is True
    assert result.has_lists is True


# ── PDF ───────────────────────────────────────────────────────────────────


def test_ingest_pdf(tmp_path: Path) -> None:
    """PDF with text should extract text content and detect headings."""
    import fitz

    pdf_file = tmp_path / "test.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Test Title", fontsize=18)
    page.insert_text((72, 120), "Body content.", fontsize=11)
    doc.save(str(pdf_file))
    doc.close()

    result = DocumentIngester().ingest(pdf_file)

    assert isinstance(result, IngestResult)
    assert "Test Title" in result.text
    assert "Body content" in result.text


# ── DOCX ──────────────────────────────────────────────────────────────────


def test_ingest_docx(tmp_path: Path) -> None:
    """DOCX with headings should extract heading structure."""
    from docx import Document

    docx_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Content.")
    doc.add_heading("Subtitle", level=2)
    doc.add_paragraph("More content.")
    doc.save(str(docx_file))

    result = DocumentIngester().ingest(docx_file)

    assert isinstance(result, IngestResult)
    assert result.title == "Title"
    assert len(result.headings) == 2
    assert result.headings[0] == Heading(level=1, text="Title")
    assert result.headings[1] == Heading(level=2, text="Subtitle")
    assert "Content" in result.text


def test_ingest_docx_with_tables(tmp_path: Path) -> None:
    """DOCX with tables should set has_tables=True."""
    from docx import Document

    docx_file = tmp_path / "table.docx"
    doc = Document()
    doc.add_heading("Doc with Table", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    doc.save(str(docx_file))

    result = DocumentIngester().ingest(docx_file)

    assert result.has_tables is True
    assert result.title == "Doc with Table"


# ── HTML ──────────────────────────────────────────────────────────────────


def test_ingest_html(tmp_path: Path) -> None:
    """HTML with h1/h2/ul should extract headings and detect lists."""
    html_file = tmp_path / "page.html"
    html_file.write_text(
        "<html><body>"
        "<h1>Page Title</h1>"
        "<p>Intro paragraph.</p>"
        "<h2>Section One</h2>"
        "<ul><li>Item A</li><li>Item B</li></ul>"
        "</body></html>",
        encoding="utf-8",
    )

    result = DocumentIngester().ingest(html_file)

    assert isinstance(result, IngestResult)
    assert result.title == "Page Title"
    assert len(result.headings) == 2
    assert result.headings[0] == Heading(level=1, text="Page Title")
    assert result.headings[1] == Heading(level=2, text="Section One")
    assert result.has_lists is True
    assert "Intro paragraph" in result.text


def test_ingest_html_with_table(tmp_path: Path) -> None:
    """HTML with table should set has_tables=True."""
    html_file = tmp_path / "table.html"
    html_file.write_text(
        "<html><body>"
        "<h1>Table Page</h1>"
        '<table><tr><td>Cell</td></tr></table>'
        "</body></html>",
        encoding="utf-8",
    )

    result = DocumentIngester().ingest(html_file)

    assert result.has_tables is True


# ── Unsupported format ────────────────────────────────────────────────────


def test_unsupported_format(tmp_path: Path) -> None:
    """Unsupported file extension should raise ValueError."""
    xyz_file = tmp_path / "data.xyz"
    xyz_file.write_text("some binary content", encoding="utf-8")

    ingester = DocumentIngester()
    raised = False
    try:
        ingester.ingest(xyz_file)
    except ValueError as exc:
        raised = True
        assert "不支持" in str(exc)

    assert raised, "Expected ValueError for unsupported format"
